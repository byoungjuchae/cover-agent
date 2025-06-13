<<<<<<< HEAD
<<<<<<< HEAD
from docx import Document
from mcp.server.fastmcp import FastMCP


mcp = FastMCP("docx_save")

@mcp.tool()
def save_docx():
    doc = Document()
    doc.save(f"{save}.docx")


if __name__ =='__main__':
    print("docx_save start")
    mcp.run(transport="stdio")
=======
from docx import Document
from mcp.server.fastmcp import FastMCP


mcp = FastMCP("docx_save")

@mcp.tool()
def save_docx():
    doc = Document()
    doc.save(f"{save}.docx")


if __name__ =='__main__':
    print("docx_save start")
    mcp.run(transport="stdio")
>>>>>>> 5b1b31de8ac8f9f71f1163adc86918c19b8e48f4
=======
from docx import Document
from mcp.server.fastmcp import FastMCP


mcp = FastMCP("docx_save")

@mcp.tool()
def save_docx():
    doc = Document()
    doc.save(f"{save}.docx")


if __name__ =='__main__':
    print("docx_save start")
    mcp.run(transport="stdio")
>>>>>>> 10602bec9b75ad2bdac436d91b87816b395cc6a1
    